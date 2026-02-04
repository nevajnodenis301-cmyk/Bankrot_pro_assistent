#!/usr/bin/env python3
"""
Create a bankruptcy petition template compatible with docxtpl (Jinja2 syntax)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins (Russian GOST standard)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.59)

# === HEADER (RIGHT-ALIGNED) ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('В {{ court_name }}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{{ court_address }}')

doc.add_paragraph()  # blank line

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('должник (заявитель):')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{{ debtor_full_name }}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('ИНН/СНИЛС: {{ debtor_inn }}/{{ debtor_snils }}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{{ debtor_address }}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('к.т. {{ debtor_phone }}')

doc.add_paragraph()  # blank line

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('кредитор:')

# Creditors list (Jinja2 for loop)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{% for creditor in creditors %}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{{ loop.index }}. {{ creditor.name }}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('ОГРН/ИНН: {{ creditor.ogrn }}/{{ creditor.inn }}')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{{ creditor.address }}')

doc.add_paragraph()  # blank line

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('{% endfor %}')

doc.add_paragraph()  # blank line

# === TITLE (CENTERED) ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ЗАЯВЛЕНИЕ')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('о признании гражданина несостоятельным (банкротом)')
run.bold = True

doc.add_paragraph()  # blank line

# === MAIN CONTENT ===
doc.add_paragraph(
    '{{ debtor_full_name }}, {{ debtor_birth_date }} года рождения, имеющий паспорт '
    '{{ debtor_passport_series }} {{ debtor_passport_number }}, выдан {{ debtor_passport_issued_by }} '
    'от {{ debtor_passport_date }}, код подразделения {{ debtor_passport_code }}, '
    'зарегистрированный и проживающий по адресу: {{ debtor_address }}, '
    'поставленный на учет в налоговом органе с присвоением ИНН/СНИЛС: {{ debtor_inn }}/{{ debtor_snils }}.'
)

doc.add_paragraph(
    '{{ debtor_surname }} {{ debtor_initials }} не зарегистрирован в качестве индивидуального предпринимателя, '
    'что подтверждается справкой № {{ ip_certificate_number }} от {{ ip_certificate_date }} года.'
)

doc.add_paragraph()

doc.add_paragraph(
    'В настоящее время общий размер кредиторской задолженности составляет '
    '{{ total_debt_rubles }} {{ total_debt_rubles_word }} {{ total_debt_kopecks }} {{ total_debt_kopecks_word }} включает в себя:'
)

doc.add_paragraph()

# Debts list
doc.add_paragraph('{% for debt in debts %}')
doc.add_paragraph('{{ loop.index }}.    задолженность перед {{ debt.creditor_name }}:')
doc.add_paragraph('- согласно официальной выписки {{ debt.source }} составляет {{ debt.amount_rubles }} {{ debt.amount_rubles_word }} {{ debt.amount_kopecks }} {{ debt.amount_kopecks_word }}.')
doc.add_paragraph()
doc.add_paragraph('{% endfor %}')

doc.add_paragraph(
    'Кредитные обязательства являются потребительскими и не связаны с осуществлением '
    'предпринимательской деятельности. Денежные средства в кредит были взяты для потребительских нужд.'
)

doc.add_paragraph(
    'Информация о счетах, открытых в банках и иных кредитных организациях приведена в форме '
    '«Опись имущества должника гражданина».'
)

doc.add_paragraph(
    'Электронные денежные средства у {{ debtor_surname }} {{ debtor_initials }} отсутствуют. '
    'Операции с такими денежными средствами за трехлетний период, предшествующий дате подачи '
    'заявления о признании банкротом, должник не совершал.'
)

# Marital status
doc.add_paragraph(
    '{% if is_married %}'
    '{{ debtor_surname }} {{ debtor_initials }} состоит в официальном браке с {{ spouse_name }}, '
    'что подтверждается свидетельством о заключении брака {{ marriage_certificate_number }} '
    'от {{ marriage_certificate_date }} года.'
    '{% elif is_divorced %}'
    '{{ debtor_surname }} {{ debtor_initials }} не состоит в официальном браке, '
    'что подтверждается свидетельством о расторжении брака {{ divorce_certificate_number }} '
    'от {{ divorce_certificate_date }} года.'
    '{% else %}'
    '{{ debtor_surname }} {{ debtor_initials }} не состоит в браке.'
    '{% endif %}'
)

# Children
doc.add_paragraph(
    '{% if has_children %}'
    'У {{ debtor_surname }} {{ debtor_initials }} на иждивении есть несовершеннолетний ребенок'
    '{% if multiple_children %} / несовершеннолетние дети{% endif %}:'
    '{% for child in children %}'
    '{{ child.child_name }}, {{ child.child_birth_date }} года рождения'
    '{% endfor %}'
    '{% else %}'
    'У {{ debtor_surname }} {{ debtor_initials }} на иждивении нет несовершеннолетних детей.'
    '{% endif %}'
)

# Employment
doc.add_paragraph(
    '{{ debtor_surname }} {{ debtor_initials }} официально не трудоустроен{% if is_self_employed %}, '
    '{{ debtor_surname }} {{ debtor_initials }} является самозанятым{% endif %}.'
)

# Property
doc.add_paragraph(
    '{% if has_real_estate %}'
    'За {{ debtor_surname }} {{ debtor_initials }} зарегистрировано недвижимое имущество: {{ real_estate_description }}'
    '{% else %}'
    'За {{ debtor_surname }} {{ debtor_initials }} не зарегистрировано недвижимое имущество.'
    '{% endif %}'
)

doc.add_paragraph(
    '{% if has_movable_property %}'
    'Движимое имущество у {{ debtor_surname }} {{ debtor_initials }} имеется, а именно {{ movable_property_description }}'
    '{% else %}'
    'Движимое имущество у {{ debtor_surname }} {{ debtor_initials }} не имеется.'
    '{% endif %}'
)

# Legal grounds
doc.add_paragraph(
    'На основании вышеизложенного, у {{ debtor_surname }} {{ debtor_initials }} имеются признаки банкротства, '
    'предусмотренные ст. 213.6 ФЗ «О несостоятельности (банкротстве)».'
)

doc.add_paragraph(
    'Прошу назначить финансового управляющего из числа членов саморегулируемой организации {{ sro_name }}'
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ПРОШУ:')
run.bold = True

doc.add_paragraph()

doc.add_paragraph('1. Признать {{ debtor_full_name }} несостоятельным (банкротом).')
doc.add_paragraph('2. Ввести в отношении {{ debtor_surname }} {{ debtor_initials }} процедуру реструктуризации долгов гражданина сроком на {{ restructuring_duration }}.')
doc.add_paragraph('3. Утвердить финансовым управляющим арбитражного управляющего из числа членов {{ sro_name }}.')

# Signature
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('{{ petition_date }}\t\t\t\t\t_____________ / {{ debtor_full_name }}')

# Save
output_path = Path(__file__).resolve().parent / "templates" / "bankruptcy_petition_template_v1_jinja2.docx"
doc.save(str(output_path))
print("✅ Template created with Jinja2 syntax!")
